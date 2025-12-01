from __future__ import annotations

import hashlib
import logging
import re
import statistics
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}


class ElementType(str, Enum):
    TABLE = "table"
    LIST = "list"
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    OTHER = "other"


@dataclass
class DocumentElement:
    text: str
    type: ElementType
    page: Optional[int] = None
    parent_id: Optional[str] = None
    heading_level: Optional[int] = None


def _clean_text(text: Optional[str]) -> str:
    """Normalize whitespace and guard against None."""
    return " ".join((text or "").split())


def _hash_id(*parts: str) -> str:
    """Stable SHA-256 hash from a sequence of strings."""
    h = hashlib.sha256()
    for part in parts:
        h.update(part.encode("utf-8", errors="ignore"))
    return h.hexdigest()


def _first_non_null(values: Sequence[Optional[int]]) -> Optional[int]:
    """Return the first non-null value in a sequence."""
    for v in values:
        if v is not None:
            return v
    return None


def _estimate_tokens(text: str) -> int:
    """Rough token estimate using whitespace splitting."""
    return max(1, len(text.split()))


def _looks_like_list_line(line: str) -> bool:
    stripped = line.strip()
    return bool(
        re.match(r"^[-*•●]\s+", stripped) or re.match(r"^\d+[.)]\s+", stripped) or stripped.endswith(":")
    )


def _strip_list_marker(line: str) -> str:
    stripped = line.strip()
    stripped = re.sub(r"^[-*•●]\s+", "", stripped)
    stripped = re.sub(r"^\d+[.)]\s+", "", stripped)
    return stripped.strip()


def _bbox_overlaps(a: Iterable[float], b: Iterable[float]) -> bool:
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    return ax0 < bx1 and ax1 > bx0 and ay0 < by1 and ay1 > by0


def _is_heading_block(text: str, avg_size: float, median_size: float) -> bool:
    if len(text) < 8 and avg_size >= median_size:
        return True
    if avg_size >= median_size + 1.5 and len(text) <= 120:
        return True
    return False


class PDFLoader:
    """Structured PDF loader using PyMuPDF heuristics."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[DocumentElement]:
        doc = fitz.open(self.file_path)
        elements: List[DocumentElement] = []

        for page_index, page in enumerate(doc, start=1):
            table_elements, table_bboxes = self._extract_tables(page, page_index)
            elements.extend(table_elements)
            elements.extend(self._extract_text_blocks(page, page_index, table_bboxes))

        return elements

    def _extract_tables(self, page: fitz.Page, page_number: int) -> tuple[List[DocumentElement], List[Iterable[float]]]:
        tables: List[DocumentElement] = []
        bboxes: List[Iterable[float]] = []
        try:
            finder = page.find_tables()
            for idx, table in enumerate(finder.tables):
                text = self._table_to_text(table)
                if not _clean_text(text):
                    continue
                tables.append(
                    DocumentElement(
                        text=text,
                        type=ElementType.TABLE,
                        page=page_number,
                        parent_id=f"tbl-{page_number}-{idx}",
                    )
                )
                if hasattr(table, "bbox"):
                    bboxes.append(table.bbox)
        except Exception:
            logger.info("PDF table extraction skipped on page %s", page_number)
        return tables, bboxes

    def _table_to_text(self, table: Any) -> str:
        if hasattr(table, "to_markdown"):
            try:
                return table.to_markdown()
            except Exception:
                pass
        if hasattr(table, "extract"):
            try:
                rows = table.extract()
                return "\n".join("\t".join(cell or "" for cell in row) for row in rows)
            except Exception:
                pass
        return ""

    def _extract_text_blocks(
        self, page: fitz.Page, page_number: int, table_bboxes: List[Iterable[float]]
    ) -> List[DocumentElement]:
        page_dict = page.get_text("dict")
        blocks = page_dict.get("blocks", [])

        span_sizes: List[float] = []
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    if span.get("text"):
                        span_sizes.append(float(span.get("size", 0)))

        median_size = statistics.median(span_sizes) if span_sizes else 12.0
        items: List[DocumentElement] = []

        for block in blocks:
            if block.get("type") != 0:
                continue
            bbox = block.get("bbox")
            if bbox and any(_bbox_overlaps(bbox, tb) for tb in table_bboxes):
                continue

            lines = block.get("lines", [])
            line_texts: List[str] = []
            line_sizes: List[float] = []
            for line in lines:
                spans = line.get("spans", [])
                text = " ".join(span.get("text", "") for span in spans)
                if text.strip():
                    line_texts.append(text)
                    line_sizes.extend([float(span.get("size", 0)) for span in spans if span.get("text")])

            text_block = "\n".join(line_texts).strip()
            if not text_block:
                continue

            avg_size = sum(line_sizes) / len(line_sizes) if line_sizes else median_size
            if _is_heading_block(text_block, avg_size, median_size):
                items.append(
                    DocumentElement(
                        text=_clean_text(text_block),
                        type=ElementType.HEADING,
                        page=page_number,
                    )
                )
                continue

            if _looks_like_list_line(text_block.splitlines()[0]):
                parent_id = f"list-{page_number}-{_hash_id(text_block)[:8]}"
                for line in text_block.splitlines():
                    if _looks_like_list_line(line):
                        clean = _strip_list_marker(line)
                        if clean:
                            items.append(
                                DocumentElement(
                                    text=clean,
                                    type=ElementType.LIST,
                                    page=page_number,
                                    parent_id=parent_id,
                                )
                            )
                    elif line.strip():
                        items.append(
                            DocumentElement(
                                text=_clean_text(line),
                                type=ElementType.PARAGRAPH,
                                page=page_number,
                                parent_id=None,
                            )
                        )
                continue

            items.append(
                DocumentElement(
                    text=_clean_text(text_block),
                    type=ElementType.PARAGRAPH,
                    page=page_number,
                )
            )
        return items


class DocxLoader:
    """Structured DOCX/DOC loader using python-docx."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[DocumentElement]:
        doc = DocxDocument(self.file_path)
        elements: List[DocumentElement] = []

        for idx, table in enumerate(doc.tables):
            text = self._table_text(table)
            if _clean_text(text):
                elements.append(
                    DocumentElement(
                        text=text,
                        type=ElementType.TABLE,
                        page=None,
                        parent_id=f"tbl-0-{idx}",
                    )
                )

        for para in doc.paragraphs:
            text = _clean_text(para.text)
            if not text:
                continue
            style = para.style.name if para.style else ""
            if style.startswith("Heading"):
                level = self._parse_heading_level(style)
                elements.append(
                    DocumentElement(
                        text=text,
                        type=ElementType.HEADING,
                        page=None,
                        heading_level=level,
                    )
                )
                continue
            if self._is_list_paragraph(para):
                num_pr = getattr(getattr(para._p, "pPr", None), "numPr", None)  # type: ignore[attr-defined]
                num_id = getattr(getattr(num_pr, "numId", None), "val", None)
                parent_id = f"list-docx-{num_id or _hash_id(text)[:8]}"
                elements.append(
                    DocumentElement(
                        text=text,
                        type=ElementType.LIST,
                        page=None,
                        parent_id=parent_id,
                    )
                )
                continue

            elements.append(DocumentElement(text=text, type=ElementType.PARAGRAPH))

        return elements

    def _table_text(self, table: Any) -> str:
        rows: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("\t".join(cells))
        return "\n".join(rows)

    def _is_list_paragraph(self, para: Any) -> bool:
        ppr = getattr(para._p, "pPr", None)  # type: ignore[attr-defined]
        num_pr = getattr(ppr, "numPr", None) if ppr is not None else None
        return num_pr is not None

    def _parse_heading_level(self, style_name: str) -> Optional[int]:
        match = re.search(r"Heading\s*(\d+)", style_name)
        return int(match.group(1)) if match else None


class MarkdownLoader:
    """Lightweight Markdown loader with heuristics for headings, lists, and tables."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[DocumentElement]:
        with open(self.file_path, "r", encoding="utf-8") as f:
            lines = f.read().splitlines()
        return self._parse_lines(lines)

    def _parse_lines(self, lines: List[str]) -> List[DocumentElement]:
        elements: List[DocumentElement] = []
        i = 0
        parent_id: Optional[str] = None

        while i < len(lines):
            line = lines[i]
            if not line.strip():
                i += 1
                parent_id = None
                continue

            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = _clean_text(line.lstrip("#").strip())
                elements.append(
                    DocumentElement(
                        text=text,
                        type=ElementType.HEADING,
                        heading_level=level,
                    )
                )
                i += 1
                parent_id = None
                continue

            # Table detection: header line with pipe + separator in next line
            if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?\s*-", lines[i + 1]):
                table_lines = [line]
                i += 1
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                elements.append(
                    DocumentElement(
                        text="\n".join(table_lines),
                        type=ElementType.TABLE,
                    )
                )
                parent_id = None
                continue

            if _looks_like_list_line(line):
                parent_id = parent_id or f"list-md-{_hash_id(str(i))[:8]}"
                elements.append(
                    DocumentElement(
                        text=_strip_list_marker(line),
                        type=ElementType.LIST,
                        parent_id=parent_id,
                    )
                )
                i += 1
                continue

            # Paragraph: gather until blank or structural line
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not _looks_like_list_line(lines[i]) and not lines[i].startswith("#"):
                para_lines.append(lines[i])
                i += 1
            elements.append(DocumentElement(text=_clean_text(" ".join(para_lines)), type=ElementType.PARAGRAPH))
            parent_id = None

        return elements


class TextLoader:
    """Plain text loader; treats blank lines as paragraph boundaries."""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[DocumentElement]:
        with open(self.file_path, "r", encoding="utf-8") as f:
            content = f.read()
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]
        return [DocumentElement(text=_clean_text(p), type=ElementType.PARAGRAPH) for p in paragraphs]


class ChunkingConfig:
    """Configuration for semantic chunking and overlaps."""

    def __init__(
        self,
        min_chunk_tokens: int = 1500,
        max_chunk_tokens: int = 2000,
        default_overlap_tokens: int = 100,
        clear_boundary_overlap_tokens: int = 0,
        sentence_split_regex: str = r"(?<=[.?!])\s+",
    ) -> None:
        if min_chunk_tokens > max_chunk_tokens:
            raise ValueError("min_chunk_tokens cannot be greater than max_chunk_tokens")
        if default_overlap_tokens < 0 or clear_boundary_overlap_tokens < 0:
            raise ValueError("overlap tokens must be non-negative")

        self.min_chunk_tokens = min_chunk_tokens
        self.max_chunk_tokens = max_chunk_tokens
        self.default_overlap_tokens = default_overlap_tokens
        self.clear_boundary_overlap_tokens = clear_boundary_overlap_tokens
        self.sentence_split_regex = sentence_split_regex


class DocumentProcessor:
    """Parse documents, preserve structure, and produce semantic chunks ready for indexing."""

    def __init__(self, chunking: Optional[ChunkingConfig] = None):
        """
        Initialize the processor with semantic chunking defaults.

        Args:
            chunking: Optional custom chunking configuration.
        """
        self.chunking = chunking or ChunkingConfig()
        self._chunker = _FallbackSemanticChunker(
            min_chunk_tokens=self.chunking.min_chunk_tokens,
            max_chunk_tokens=self.chunking.max_chunk_tokens,
            sentence_split_regex=self.chunking.sentence_split_regex,
        )

    def load_document(self, file_path: str) -> List[DocumentElement]:
        """
        Load a document with format-specific loaders, preserving element types.

        Raises:
            ValueError: if the file extension is not supported or parsing fails.
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        if ext == ".pdf":
            loader = PDFLoader(file_path)
        elif ext in (".docx", ".doc"):
            loader = DocxLoader(file_path)
        elif ext == ".md":
            loader = MarkdownLoader(file_path)
        else:
            loader = TextLoader(file_path)

        try:
            return [el for el in loader.load() if _clean_text(el.text)]
        except Exception as exc:
            logger.exception("Failed to load document %s", file_path)
            raise ValueError(f"Failed to parse document: {exc}") from exc

    def process_document(
        self,
        file_path: str,
        document_id: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Document]:
        """
        Full ingestion pipeline:
        1) Load with format-specific loader.
        2) Preprocess by element type (tables, lists, headings, paragraphs).
        3) Semantic chunking for paragraph content.
        4) Enrich metadata and generate stable chunk IDs.
        """
        elements = self.load_document(file_path)
        if not elements:
            return []

        docs: List[Document] = []
        list_buffer: List[DocumentElement] = []
        paragraph_buffer: List[DocumentElement] = []
        chunk_index = 0
        current_heading: Optional[str] = None

        def flush_list_buffer() -> None:
            """Group contiguous list items under the same parent block."""
            nonlocal chunk_index
            if not list_buffer:
                return
            parent_id = list_buffer[0].parent_id
            text_lines = [f"- {_clean_text(item.text)}" for item in list_buffer if _clean_text(item.text)]
            page_numbers = [item.page for item in list_buffer]
            list_text = "\n".join(text_lines)
            if not list_text:
                list_buffer.clear()
                return

            md = self._build_metadata(
                metadata=metadata,
                document_id=document_id,
                source=file_path,
                page=_first_non_null(page_numbers),
                element_type=ElementType.LIST.value,
                heading=current_heading,
                extra={"list_parent_id": parent_id},
            )
            docs.append(self._to_document(list_text, md, chunk_index, document_id))
            chunk_index += 1
            list_buffer.clear()

        def flush_paragraph_buffer() -> None:
            """Send buffered paragraphs to the semantic chunker."""
            nonlocal chunk_index
            if not paragraph_buffer:
                return
            page = _first_non_null([p.page for p in paragraph_buffer])
            text = "\n\n".join(_clean_text(p.text) for p in paragraph_buffer if _clean_text(p.text))
            paragraph_buffer.clear()
            if not text.strip():
                return

            base_md = self._build_metadata(
                metadata=metadata,
                document_id=document_id,
                source=file_path,
                page=page,
                element_type=ElementType.PARAGRAPH.value,
                heading=current_heading,
            )
            paragraph_docs = self._semantic_chunks(text, base_md, chunk_index, document_id)
            docs.extend(paragraph_docs)
            chunk_index += len(paragraph_docs)

        for element in elements:
            if element.type == ElementType.LIST:
                flush_paragraph_buffer()
                if list_buffer and list_buffer[-1].parent_id != element.parent_id:
                    flush_list_buffer()
                list_buffer.append(element)
                continue

            flush_list_buffer()

            if element.type == ElementType.TABLE:
                flush_paragraph_buffer()
                table_text = _clean_text(element.text)
                if not table_text:
                    continue
                md = self._build_metadata(
                    metadata=metadata,
                    document_id=document_id,
                    source=file_path,
                    page=element.page,
                    element_type=ElementType.TABLE.value,
                    heading=current_heading,
                )
                docs.append(self._to_document(table_text, md, chunk_index, document_id))
                chunk_index += 1
                continue

            if element.type == ElementType.HEADING:
                flush_paragraph_buffer()
                heading_text = _clean_text(element.text)
                if heading_text:
                    md = self._build_metadata(
                        metadata=metadata,
                        document_id=document_id,
                        source=file_path,
                        page=element.page,
                        element_type=ElementType.HEADING.value,
                        heading=None,
                        extra={"heading_level": element.heading_level},
                    )
                    docs.append(self._to_document(heading_text, md, chunk_index, document_id))
                    chunk_index += 1
                    current_heading = heading_text
                continue

            if element.type == ElementType.PARAGRAPH:
                paragraph_buffer.append(element)
                continue

            # Fallback
            paragraph_buffer.append(
                DocumentElement(text=element.text, type=ElementType.PARAGRAPH, page=element.page)
            )

        flush_list_buffer()
        flush_paragraph_buffer()

        cleaned = [d for d in docs if d.page_content.strip()]
        logger.info("Chunked document %s into %d chunks", file_path, len(cleaned))
        return cleaned

    def _build_metadata(
        self,
        metadata: Optional[Dict[str, Any]],
        document_id: str,
        source: str,
        page: Optional[int],
        element_type: str,
        heading: Optional[str],
        extra: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Compose metadata shared across all chunk types."""
        base = dict(metadata or {})
        base.update(
            {
                "document_id": document_id,
                "source": source,
                "page": page,
                "element_type": element_type,
                "heading": heading,
            }
        )
        if extra:
            base.update(extra)
        return base

    def _semantic_chunks(
        self,
        text: str,
        metadata: Dict[str, Any],
        start_index: int,
        document_id: str,
    ) -> List[Document]:
        """Run semantic chunking, enforce size windows, and attach IDs."""
        raw_chunks = self._chunker.split_text(text)
        sized_chunks = self._enforce_size(raw_chunks)
        overlapped = self._apply_dynamic_overlap(sized_chunks)

        documents: List[Document] = []
        for offset, chunk in enumerate(overlapped):
            md = dict(metadata)
            md["chunk_index"] = start_index + offset
            md["chunk_id"] = _hash_id(document_id, str(md.get("page")), str(md["chunk_index"]), chunk)
            documents.append(Document(page_content=chunk, metadata=md))
        return documents

    def _enforce_size(self, chunks: List[str]) -> List[str]:
        """Merge or split semantic chunks to stay within the configured token window."""
        merged: List[str] = []
        buffer: List[str] = []
        buffer_tokens = 0

        for chunk in chunks:
            chunk_tokens = _estimate_tokens(chunk)

            if chunk_tokens >= self.chunking.max_chunk_tokens:
                # Flush buffer before handling oversized chunk
                if buffer:
                    merged.append("\n\n".join(buffer))
                    buffer = []
                    buffer_tokens = 0
                merged.extend(self._split_oversize(chunk))
                continue

            if buffer_tokens + chunk_tokens > self.chunking.max_chunk_tokens and buffer:
                merged.append("\n\n".join(buffer))
                buffer = [chunk]
                buffer_tokens = chunk_tokens
                continue

            buffer.append(chunk)
            buffer_tokens += chunk_tokens

            if buffer_tokens >= self.chunking.min_chunk_tokens:
                merged.append("\n\n".join(buffer))
                buffer = []
                buffer_tokens = 0

        if buffer:
            merged.append("\n\n".join(buffer))

        return merged

    def _split_oversize(self, text: str) -> List[str]:
        """Split a single oversized chunk while honoring clear boundaries."""
        words = text.split()
        slices: List[str] = []
        start = 0
        max_tokens = self.chunking.max_chunk_tokens

        while start < len(words):
            end = min(len(words), start + max_tokens)
            slices.append(" ".join(words[start:end]))
            start = end

        return slices

    def _apply_dynamic_overlap(self, chunks: List[str]) -> List[str]:
        """Apply 0 or 100 token overlap depending on boundary clarity."""
        if not chunks:
            return []

        overlapped: List[str] = [chunks[0]]
        previous_raw = chunks[0]

        for chunk in chunks[1:]:
            overlap_tokens = (
                self.chunking.clear_boundary_overlap_tokens
                if self._has_clear_boundary(previous_raw) or self._has_clear_boundary(chunk)
                else self.chunking.default_overlap_tokens
            )
            if overlap_tokens > 0:
                tail = " ".join(previous_raw.split()[-overlap_tokens:])
                chunk_with_overlap = f"{tail}\n\n{chunk}" if tail else chunk
            else:
                chunk_with_overlap = chunk
            overlapped.append(chunk_with_overlap)
            previous_raw = chunk

        return overlapped

    def _has_clear_boundary(self, text: str) -> bool:
        """Detect if a chunk is naturally separated from neighbors."""
        if "\n\n" in text:
            return True
        boundary_markers = ("##", "--", "===", "Section", "Chapter")
        return any(marker in text for marker in boundary_markers)

    def _to_document(
        self, text: str, metadata: Dict[str, Any], chunk_index: int, document_id: str
    ) -> Document:
        """Create a LangChain Document with stable chunk metadata."""
        md = dict(metadata)
        md["chunk_index"] = chunk_index
        md["chunk_id"] = _hash_id(document_id, str(md.get("page")), str(chunk_index), text)
        return Document(page_content=text, metadata=md)


class _FallbackSemanticChunker:
    """
    Minimal semantic-ish chunker used when SemanticChunker is unavailable.

    Splits by sentences, then merges to stay within min/max token targets.
    """

    def __init__(
        self,
        min_chunk_tokens: int = 1500,
        max_chunk_tokens: int = 2000,
        sentence_split_regex: str = r"(?<=[.?!])\s+",
    ):
        self.min_chunk_tokens = min_chunk_tokens
        self.max_chunk_tokens = max_chunk_tokens
        self.sentence_split_regex = re.compile(sentence_split_regex)

    def split_text(self, text: str) -> List[str]:
        sentences = [s.strip() for s in self.sentence_split_regex.split(text) if s.strip()]
        if not sentences:
            return [text]

        chunks: List[str] = []
        buffer: List[str] = []
        tokens = 0

        for sent in sentences:
            sent_tokens = _estimate_tokens(sent)
            if tokens + sent_tokens > self.max_chunk_tokens and buffer:
                chunks.append(" ".join(buffer))
                buffer = [sent]
                tokens = sent_tokens
                continue

            buffer.append(sent)
            tokens += sent_tokens
            if tokens >= self.min_chunk_tokens:
                chunks.append(" ".join(buffer))
                buffer = []
                tokens = 0

        if buffer:
            chunks.append(" ".join(buffer))

        return chunks
